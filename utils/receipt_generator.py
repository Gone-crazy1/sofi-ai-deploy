"""
Beautiful Receipt Generator for Sofi AI
Creates professional HTML and PDF receipts for transactions
"""

import os
import logging
from typing import Dict, Any, Optional
from datetime import datetime
from jinja2 import Template

logger = logging.getLogger(__name__)

# Try to import PDF generation libraries
try:
    import weasyprint
    PDF_AVAILABLE = True
    logger.info("✅ WeasyPrint available - PDF generation enabled")
except ImportError as e:
    PDF_AVAILABLE = False
    logger.warning(f"WeasyPrint not available - trying alternative PDF methods: {e}")
except Exception as e:
    PDF_AVAILABLE = False
    logger.warning(f"WeasyPrint import error - trying alternative PDF methods: {e}")

# Try alternative PDF generation
ALTERNATIVE_PDF_AVAILABLE = False
try:
    # Try reportlab for simple PDF generation
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import inch
    ALTERNATIVE_PDF_AVAILABLE = True
    logger.info("✅ ReportLab available - PDF generation enabled")
except ImportError as e:
    logger.warning(f"ReportLab not available: {e}")

# Try to import image generation libraries
IMAGE_GENERATION_AVAILABLE = False
try:
    from PIL import Image, ImageDraw, ImageFont
    import io
    import base64
    IMAGE_GENERATION_AVAILABLE = True
    logger.info("✅ PIL available - Image generation enabled")
except ImportError as e:
    logger.warning(f"PIL not available for image generation: {e}")

# Try to import HTML to image conversion
HTML_TO_IMAGE_AVAILABLE = False
try:
    from html2image import Html2Image
    HTML_TO_IMAGE_AVAILABLE = True
    logger.info("✅ html2image available - HTML to image conversion enabled")
except ImportError as e:
    logger.warning(f"html2image not available: {e}")

# Try selenium as backup for HTML to image
SELENIUM_AVAILABLE = False
try:
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from webdriver_manager.chrome import ChromeDriverManager
    SELENIUM_AVAILABLE = True
    logger.info("✅ Selenium available - HTML to image conversion enabled")
except ImportError as e:
    logger.warning(f"Selenium not available: {e}")

class SofiReceiptGenerator:
    """Generate beautiful receipts for Sofi AI transactions"""
    
    def __init__(self):
        self.template_html = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Sofi AI - Transaction Receipt</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif;
            background: #f8f9fa;
            padding: 15px;
            width: 350px;
            height: auto;
        }
        
        .receipt-container {
            background: white;
            border-radius: 12px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            overflow: hidden;
            border: 1px solid #e9ecef;
            position: relative;
        }
        
        /* Watermark */
        .watermark {
            position: absolute;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%) rotate(-45deg);
            font-size: 60px;
            color: rgba(76, 175, 80, 0.08);
            font-weight: bold;
            pointer-events: none;
            z-index: 1;
            letter-spacing: 2px;
        }
        
        .receipt-header {
            background: linear-gradient(135deg, #4CAF50 0%, #45a049 100%);
            color: white;
            padding: 25px 20px;
            text-align: center;
            position: relative;
            z-index: 2;
        }
        
        .logo {
            font-size: 20px;
            font-weight: 700;
            margin-bottom: 8px;
        }
        
        .amount-display {
            font-size: 32px;
            font-weight: 700;
            margin: 12px 0;
        }
        
        .status {
            font-size: 16px;
            opacity: 0.95;
            font-weight: 500;
        }
        
        .timestamp {
            font-size: 13px;
            opacity: 0.8;
            margin-top: 8px;
        }
        
        .receipt-body {
            padding: 20px;
            position: relative;
            z-index: 2;
        }
        
        .info-section {
            margin-bottom: 18px;
            padding-bottom: 18px;
            border-bottom: 1px solid #f1f3f4;
        }
        
        .info-section:last-child {
            border-bottom: none;
            margin-bottom: 0;
        }
        
        .section-title {
            font-size: 12px;
            color: #666;
            margin-bottom: 8px;
            text-transform: uppercase;
            font-weight: 600;
            letter-spacing: 0.5px;
        }
        
        .main-info {
            font-size: 16px;
            color: #333;
            font-weight: 700;
            margin-bottom: 4px;
            text-transform: uppercase;
        }
        
        .sub-info {
            font-size: 13px;
            color: #666;
            line-height: 1.4;
        }
        
        .amount-grid {
            display: grid;
            grid-template-columns: 1fr auto;
            gap: 8px;
        }
        
        .amount-row {
            display: contents;
            font-size: 13px;
        }
        
        .amount-label {
            color: #666;
            font-weight: 500;
        }
        
        .amount-value {
            color: #333;
            font-weight: 600;
            text-align: right;
        }
        
        .total-row {
            border-top: 1px solid #eee;
            padding-top: 8px;
            margin-top: 8px;
            font-weight: 700;
            font-size: 14px;
        }
        
        .reference {
            font-family: 'Courier New', monospace;
            font-size: 11px;
            color: #666;
            background: #f8f9fa;
            padding: 10px;
            border-radius: 6px;
            text-align: center;
            word-break: break-all;
            line-height: 1.3;
        }
        
        .footer-message {
            text-align: center;
            font-size: 12px;
            color: #4CAF50;
            margin-top: 15px;
            font-weight: 500;
        }
        
        .keep-record {
            text-align: center;
            font-size: 11px;
            color: #999;
            margin-top: 5px;
        }
    </style>
</head>
<body>
    <div class="receipt-container">
        <div class="watermark">Sofi AI</div>
        
        <div class="receipt-header">
            <div class="logo">Sofi AI</div>
            <div class="amount-display">₦{{ "{:,.0f}".format(amount) }}</div>
            <div class="status">Successful</div>
            <div class="timestamp">{{ transaction_time }}</div>
        </div>
        
        <div class="receipt-body">
            <div class="info-section">
                <div class="section-title">👤 Recipient</div>
                <div class="main-info">{{ recipient_name }}</div>
                <div class="sub-info">{{ bank_name.split('(')[0].strip() }} ({{ bank_name.split('(')[1].replace(')', '') if '(' in bank_name else bank_name }}) | {{ account_number }}</div>
            </div>
            
            <div class="info-section">
                <div class="amount-grid">
                    <div class="amount-row">
                        <span class="amount-label">💰 Amount Sent</span>
                        <span class="amount-value">₦{{ "{:,.0f}".format(amount) }}</span>
                    </div>
                    <div class="amount-row">
                        <span class="amount-label">💸 Transfer Fee</span>
                        <span class="amount-value">₦{{ "{:,.0f}".format(fee) }}</span>
                    </div>
                    <div class="amount-row total-row">
                        <span class="amount-label">💵 Total Charged</span>
                        <span class="amount-value">₦{{ "{:,.0f}".format(total_charged) }}</span>
                    </div>
                    <div class="amount-row">
                        <span class="amount-label">💳 New Balance</span>
                        <span class="amount-value">₦{{ "{:,.0f}".format(new_balance) }}</span>
                    </div>
                </div>
            </div>
            
            <div class="info-section">
                <div class="section-title">🧾 Reference</div>
                <div class="reference">{{ reference[:20] + '...' if reference|length > 20 else reference }}</div>
            </div>
            
            <div class="footer-message">Thank you for using Sofi AI! 💚</div>
            <div class="keep-record">Keep this receipt for your records 📄</div>
        </div>
    </div>
</body>
</html>
"""
    
    def generate_html_receipt(self, transaction_data: Dict[str, Any]) -> str:
        """Generate HTML receipt"""
        try:
            template = Template(self.template_html)
            
            # Format transaction time
            if isinstance(transaction_data.get('transaction_time'), str):
                transaction_time = transaction_data['transaction_time']
            else:
                transaction_time = datetime.now().strftime("%d/%m/%Y %I:%M %p")
            
            # Prepare template data
            template_data = {
                'amount': float(transaction_data.get('amount', 0)),
                'fee': float(transaction_data.get('fee', 0)),
                'total_charged': float(transaction_data.get('total_charged', 0)),
                'new_balance': float(transaction_data.get('new_balance', 0)),
                'recipient_name': transaction_data.get('recipient_name', 'Unknown Recipient'),
                'bank_name': transaction_data.get('bank_name', 'Unknown Bank'),
                'account_number': transaction_data.get('account_number', ''),
                'reference': transaction_data.get('reference', ''),
                'transaction_id': transaction_data.get('transaction_id', ''),
                'transaction_time': transaction_time,
                'narration': transaction_data.get('narration', '')
            }
            
            html_content = template.render(**template_data)
            return html_content
            
        except Exception as e:
            logger.error(f"Error generating HTML receipt: {e}")
            return None
    
    def generate_pdf_receipt(self, transaction_data: Dict[str, Any], output_path: Optional[str] = None) -> Optional[str]:
        """Generate PDF receipt using multiple methods"""
        try:
            # Try WeasyPrint first (if available)
            if PDF_AVAILABLE:
                try:
                    html_content = self.generate_html_receipt(transaction_data)
                    if html_content:
                        if not output_path:
                            timestamp = int(datetime.now().timestamp())
                            output_path = f"sofi_receipt_{timestamp}.pdf"
                        
                        weasyprint.HTML(string=html_content).write_pdf(output_path)
                        logger.info(f"PDF receipt generated with WeasyPrint: {output_path}")
                        return output_path
                except Exception as e:
                    logger.warning(f"WeasyPrint PDF generation failed: {e}")
            
            # Fallback to ReportLab (more reliable on Windows)
            if ALTERNATIVE_PDF_AVAILABLE:
                try:
                    from reportlab.pdfgen import canvas
                    from reportlab.lib.pagesizes import letter, A4
                    from reportlab.lib.colors import HexColor
                    from reportlab.lib.units import inch
                    
                    if not output_path:
                        timestamp = int(datetime.now().timestamp())
                        output_path = f"sofi_receipt_{timestamp}.pdf"
                    
                    # Create PDF with ReportLab
                    c = canvas.Canvas(output_path, pagesize=letter)
                    width, height = letter
                    
                    # Header
                    c.setFillColor(HexColor('#4CAF50'))
                    c.rect(0, height - 2*inch, width, 2*inch, fill=True, stroke=False)
                    
                    # Title
                    c.setFillColor(HexColor('#FFFFFF'))
                    c.setFont("Helvetica-Bold", 24)
                    c.drawString(width/2 - 50, height - 1*inch, "💳 Sofi AI")
                    c.setFont("Helvetica", 12)
                    c.drawString(width/2 - 60, height - 1.3*inch, "Transaction Receipt")
                    c.drawString(width/2 - 30, height - 1.6*inch, "✅ Successful")
                    
                    # Transaction details
                    c.setFillColor(HexColor('#000000'))
                    y_pos = height - 3*inch
                    
                    details = [
                        ("Amount Sent", f"₦{float(transaction_data.get('amount', 0)):,.2f}"),
                        ("Transfer Fee", f"₦{float(transaction_data.get('fee', 0)):,.2f}"),
                        ("Total Charged", f"₦{float(transaction_data.get('total_charged', 0)):,.2f}"),
                        ("New Balance", f"₦{float(transaction_data.get('new_balance', 0)):,.2f}"),
                        ("Recipient", transaction_data.get('recipient_name', 'Unknown')),
                        ("Bank", transaction_data.get('bank_name', 'Unknown Bank')),
                        ("Account", transaction_data.get('account_number', '')),
                        ("Reference", transaction_data.get('reference', '')),
                        ("Transaction ID", transaction_data.get('transaction_id', '')),
                        ("Time", transaction_data.get('transaction_time', datetime.now().strftime('%d/%m/%Y %I:%M %p')))
                    ]
                    
                    c.setFont("Helvetica-Bold", 14)
                    c.drawString(1*inch, y_pos, "Transaction Details")
                    y_pos -= 0.5*inch
                    
                    c.setFont("Helvetica", 10)
                    for label, value in details:
                        c.drawString(1*inch, y_pos, f"{label}:")
                        c.drawRightString(width - 1*inch, y_pos, str(value))
                        y_pos -= 0.3*inch
                    
                    # Footer
                    c.drawString(width/2 - 100, 1*inch, "Thank you for using Sofi AI!")
                    c.drawString(width/2 - 120, 0.7*inch, "Keep this receipt for your records")
                    
                    c.save()
                    logger.info(f"PDF receipt generated with ReportLab: {output_path}")
                    return output_path
                    
                except Exception as e:
                    logger.error(f"ReportLab PDF generation failed: {e}")
            
            logger.warning("No PDF generation method available")
            return None
            
        except Exception as e:
            logger.error(f"Error generating PDF receipt: {e}")
            return None
    
    def generate_image_receipt(self, transaction_data: Dict[str, Any], output_path: Optional[str] = None) -> Optional[str]:
        """Generate receipt as image using HTML to image conversion"""
        try:
            # Generate HTML first
            html_content = self.generate_html_receipt(transaction_data)
            if not html_content:
                logger.error("Failed to generate HTML content for image")
                return None
            
            if not output_path:
                timestamp = int(datetime.now().timestamp())
                output_path = f"sofi_receipt_{timestamp}.png"
            
            # Method 1: Try html2image first
            if HTML_TO_IMAGE_AVAILABLE:
                try:
                    hti = Html2Image(output_path=os.path.dirname(output_path) or ".")
                    hti.screenshot(
                        html_str=html_content,
                        css_str="",
                        save_as=os.path.basename(output_path),
                        size=(350, 500)
                    )
                    
                    # Check if file was created
                    if os.path.exists(output_path):
                        logger.info(f"Receipt image generated with html2image: {output_path}")
                        return output_path
                    else:
                        logger.warning("html2image did not create the expected file")
                        
                except Exception as e:
                    logger.warning(f"html2image failed: {e}")
            
            # Method 2: Try Selenium as backup
            if SELENIUM_AVAILABLE:
                try:
                    # Setup Chrome options for headless mode
                    chrome_options = Options()
                    chrome_options.add_argument("--headless")
                    chrome_options.add_argument("--no-sandbox")
                    chrome_options.add_argument("--disable-dev-shm-usage")
                    chrome_options.add_argument("--window-size=350,500")
                    
                    # Create driver
                    driver = webdriver.Chrome(
                        service=webdriver.chrome.service.Service(ChromeDriverManager().install()),
                        options=chrome_options
                    )
                    
                    # Create temporary HTML file
                    import tempfile
                    with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
                        f.write(html_content)
                        temp_html_path = f.name
                    
                    # Load HTML and take screenshot
                    driver.get(f"file://{temp_html_path}")
                    driver.save_screenshot(output_path)
                    driver.quit()
                    
                    # Clean up temp file
                    try:
                        os.remove(temp_html_path)
                    except:
                        pass
                    
                    if os.path.exists(output_path):
                        logger.info(f"Receipt image generated with Selenium: {output_path}")
                        return output_path
                        
                except Exception as e:
                    logger.warning(f"Selenium screenshot failed: {e}")
            
            # Method 3: Fallback to PIL-based simple image generation
            if IMAGE_GENERATION_AVAILABLE:
                try:
                    # Create a simple image with PIL
                    from PIL import Image, ImageDraw, ImageFont
                    
                    # Create image
                    width, height = 350, 500
                    img = Image.new('RGB', (width, height), color='white')
                    draw = ImageDraw.Draw(img)
                    
                    # Try to use default font
                    try:
                        font_large = ImageFont.truetype("arial.ttf", 20)
                        font_medium = ImageFont.truetype("arial.ttf", 14)
                        font_small = ImageFont.truetype("arial.ttf", 12)
                    except:
                        font_large = ImageFont.load_default()
                        font_medium = ImageFont.load_default()
                        font_small = ImageFont.load_default()
                    
                    # Draw header background
                    draw.rectangle([0, 0, width, 100], fill='#4CAF50')
                    
                    # Draw header text
                    draw.text((width//2 - 40, 15), "Sofi AI", fill='white', font=font_large)
                    draw.text((width//2 - 80, 40), f"₦{float(transaction_data.get('amount', 0)):,.0f}", fill='white', font=font_large)
                    draw.text((width//2 - 60, 65), "Transfer Successful", fill='white', font=font_medium)
                    
                    # Draw watermark
                    draw.text((width//2 - 30, height//2), "Sofi AI", fill='#E8F5E8', font=font_large)
                    
                    # Draw details
                    y_pos = 120
                    details = [
                        ("Recipient", transaction_data.get('recipient_name', 'Unknown')),
                        ("Bank", transaction_data.get('bank_name', 'Unknown Bank').split('(')[0].strip()),
                        ("Account", transaction_data.get('account_number', '')),
                        ("Amount", f"₦{float(transaction_data.get('amount', 0)):,.0f}"),
                        ("Fee", f"₦{float(transaction_data.get('fee', 0)):,.0f}"),
                        ("Total", f"₦{float(transaction_data.get('total_charged', 0)):,.0f}"),
                        ("Reference", transaction_data.get('reference', '')),
                        ("Time", transaction_data.get('transaction_time', datetime.now().strftime('%d/%m/%Y %I:%M %p')))
                    ]
                    
                    for label, value in details:
                        draw.text((20, y_pos), f"{label}:", fill='#666666', font=font_small)
                        draw.text((20, y_pos + 15), str(value)[:35], fill='#333333', font=font_medium)
                        y_pos += 45
                    
                    # Save image
                    img.save(output_path, 'PNG')
                    logger.info(f"Receipt image generated with PIL: {output_path}")
                    return output_path
                    
                except Exception as e:
                    logger.warning(f"PIL image generation failed: {e}")
            
            logger.warning("No image generation method available")
            return None
            
        except Exception as e:
            logger.error(f"Error generating image receipt: {e}")
            return None
    
    def generate_word_doc_receipt(self, transaction_data: Dict[str, Any], output_path: Optional[str] = None) -> Optional[str]:
        """Generate Word document receipt"""
        try:
            # Try to import python-docx
            try:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                logger.warning("python-docx not available - Word document generation disabled")
                return None
            
            # Create document
            doc = Document()
            
            # Add header
            header = doc.add_heading('SOFI AI - TRANSACTION RECEIPT', 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add status
            status_para = doc.add_paragraph('✅ TRANSFER SUCCESSFUL')
            status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph('')  # Empty line
            
            # Add transaction details
            details_table = doc.add_table(rows=8, cols=2)
            details_table.style = 'Light Grid Accent 1'
            
            # Fill table with data
            details = [
                ('Amount Sent', f"₦{float(transaction_data.get('amount', 0)):,.2f}"),
                ('Transfer Fee', f"₦{float(transaction_data.get('fee', 0)):,.2f}"),
                ('Total Charged', f"₦{float(transaction_data.get('total_charged', 0)):,.2f}"),
                ('New Balance', f"₦{float(transaction_data.get('new_balance', 0)):,.2f}"),
                ('Recipient', transaction_data.get('recipient_name', 'Unknown')),
                ('Bank', transaction_data.get('bank_name', 'Unknown Bank')),
                ('Account Number', transaction_data.get('account_number', '')),
                ('Reference', transaction_data.get('reference', ''))
            ]
            
            for i, (label, value) in enumerate(details):
                details_table.cell(i, 0).text = label
                details_table.cell(i, 1).text = value
            
            # Add timestamp
            doc.add_paragraph('')
            time_para = doc.add_paragraph(f"Transaction Time: {transaction_data.get('transaction_time', datetime.now().strftime('%d/%m/%Y %I:%M %p'))}")
            time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add footer
            doc.add_paragraph('')
            footer = doc.add_paragraph('Thank you for using Sofi AI!')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            if not output_path:
                timestamp = int(datetime.now().timestamp())
                output_path = f"/tmp/sofi_receipt_{timestamp}.docx"
            
            doc.save(output_path)
            
            logger.info(f"Word document receipt generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating Word document receipt: {e}")
            return None
    
    def generate_telegram_receipt(self, transaction_data: Dict[str, Any]) -> str:
        """Generate formatted receipt for Telegram"""
        try:
            amount = float(transaction_data.get('amount', 0))
            fee = float(transaction_data.get('fee', 0))
            total_charged = float(transaction_data.get('total_charged', 0))
            new_balance = float(transaction_data.get('new_balance', 0))
            
            receipt = f"""🎉 *TRANSFER SUCCESSFUL!* 🎉

━━━━━━━━━━━━━━━━━━━━━━━━━
💳 *SOFI AI RECEIPT*
━━━━━━━━━━━━━━━━━━━━━━━━━

💰 *Amount Sent:* ₦{amount:,.0f}
💸 *Transfer Fee:* ₦{fee:,.0f}
💵 *Total Charged:* ₦{total_charged:,.0f}
💳 *New Balance:* ₦{new_balance:,.0f}

👤 *Recipient:* {transaction_data.get('recipient_name', 'Unknown')}
🏦 *Bank:* {transaction_data.get('bank_name', 'Unknown Bank').split('(')[0].strip()}
📱 *Account:* {transaction_data.get('account_number', '')}

🧾 *Reference:* `{transaction_data.get('reference', '')}`
🆔 *Transaction ID:* `{transaction_data.get('transaction_id', '')}`
🕐 *Time:* {transaction_data.get('transaction_time', datetime.now().strftime('%d/%m/%Y %I:%M %p'))}

━━━━━━━━━━━━━━━━━━━━━━━━━
Thank you for using Sofi AI! 💚
Keep this receipt for your records 📄"""
            return receipt.strip()
            
        except Exception as e:
            logger.error(f"Error generating Telegram receipt: {e}")
            return "Receipt generation failed"

def create_transaction_receipt(transaction_data: Dict[str, Any], format_type: str = "telegram") -> str:
    """
    Create a transaction receipt
    
    Args:
        transaction_data: Transaction details
        format_type: "telegram", "html", or "pdf"
    
    Returns:
        Formatted receipt string
    """
    generator = SofiReceiptGenerator()
    
    if format_type == "html":
        return generator.generate_html_receipt(transaction_data)
    elif format_type == "telegram":
        return generator.generate_telegram_receipt(transaction_data)
    else:
        return generator.generate_telegram_receipt(transaction_data)
