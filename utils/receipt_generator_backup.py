"""
Beautiful Receipt Generator for Sofi AI
Creates professional HTML and PDF receipts for transactions
"""

import os
import logging
from typing import Dict, Any, Optional
from datetime import datetime
from jinja2 import Template

# Try to import weasyprint for PDF generation
try:
    import weasyprint
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("weasypdef create_transaction_receipt(transaction_data: Dict[str, Any], format_type: str = "telegram") -> str:
    """
    Create a transaction receipt
    
    Args:
        transaction_data: Transaction details
        format_type: "telegram", "html", or "pdf"
    
    Returns:
        Formatted receipt string
    """
    generator = SofiReceiptGenerator()
    
    if format_type == "html":
        return generator.generate_html_receipt(transaction_data)
    elif format_type == "telegram":
        return generator.generate_telegram_receipt(transaction_data)
    else:
        return generator.generate_telegram_receipt(transaction_data)e - PDF generation disabled")

logger = logging.getLogger(__name__)

class SofiReceiptGenerator:
    """Generate beautiful receipts for Sofi AI transactions"""
    
    def __init__(self):
        self.template_html = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Sofi AI - Transaction Receipt</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }
        
        .receipt-container {
            max-width: 400px;
            margin: 0 auto;
            background: white;
            border-radius: 15px;
            box-shadow: 0 20px 40px rgba(0,0,0,0.1);
            overflow: hidden;
        }
        
        .receipt-header {
            background: linear-gradient(135deg, #4CAF50 0%, #45a049 100%);
            color: white;
            padding: 25px 20px;
            text-align: center;
        }
        
        .logo {
            font-size: 28px;
            font-weight: bold;
            margin-bottom: 5px;
        }
        
        .tagline {
            font-size: 14px;
            opacity: 0.9;
        }
        
        .status-badge {
            background: #fff;
            color: #4CAF50;
            padding: 8px 16px;
            border-radius: 20px;
            font-weight: bold;
            font-size: 14px;
            margin-top: 15px;
            display: inline-block;
        }
        
        .receipt-body {
            padding: 25px 20px;
        }
        
        .transaction-details {
            margin-bottom: 20px;
        }
        
        .detail-row {
            display: flex;
            justify-content: space-between;
            align-items: center;
            padding: 12px 0;
            border-bottom: 1px solid #f0f0f0;
        }
        
        .detail-row:last-child {
            border-bottom: none;
        }
        
        .detail-label {
            color: #666;
            font-size: 14px;
        }
        
        .detail-value {
            font-weight: bold;
            color: #333;
            text-align: right;
            font-size: 14px;
        }
        
        .amount-highlight {
            color: #4CAF50;
            font-size: 18px;
        }
        
        .recipient-info {
            background: #f8f9fa;
            padding: 15px;
            border-radius: 10px;
            margin: 20px 0;
        }
        
        .recipient-name {
            font-weight: bold;
            color: #333;
            margin-bottom: 5px;
        }
        
        .recipient-account {
            color: #666;
            font-size: 14px;
        }
        
        .receipt-footer {
            background: #f8f9fa;
            padding: 20px;
            text-align: center;
            border-top: 1px solid #eee;
        }
        
        .footer-text {
            color: #666;
            font-size: 12px;
            line-height: 1.5;
        }
        
        .qr-placeholder {
            width: 80px;
            height: 80px;
            background: #e0e0e0;
            border-radius: 8px;
            margin: 15px auto;
            display: flex;
            align-items: center;
            justify-content: center;
            color: #999;
            font-size: 12px;
        }
        
        @media print {
            body {
                background: white;
                padding: 0;
            }
            
            .receipt-container {
                box-shadow: none;
                max-width: none;
            }
        }
    </style>
</head>
<body>
    <div class="receipt-container">
        <div class="receipt-header">
            <div class="logo">💳 SOFI AI</div>
            <div class="tagline">Your Smart Banking Assistant</div>
            <div class="status-badge">✅ SUCCESSFUL</div>
        </div>
        
        <div class="receipt-body">
            <div class="transaction-details">
                <div class="detail-row">
                    <span class="detail-label">Amount Sent</span>
                    <span class="detail-value amount-highlight">₦{{ "%.2f"|format(amount) }}</span>
                </div>
                
                <div class="detail-row">
                    <span class="detail-label">Transfer Fee</span>
                    <span class="detail-value">₦{{ "%.2f"|format(fee) }}</span>
                </div>
                
                <div class="detail-row">
                    <span class="detail-label">Total Charged</span>
                    <span class="detail-value amount-highlight">₦{{ "%.2f"|format(total_charged) }}</span>
                </div>
                
                <div class="detail-row">
                    <span class="detail-label">New Balance</span>
                    <span class="detail-value">₦{{ "%.2f"|format(new_balance) }}</span>
                </div>
            </div>
            
            <div class="recipient-info">
                <div class="recipient-name">{{ recipient_name }}</div>
                <div class="recipient-account">{{ bank_name }} • {{ account_number }}</div>
            </div>
            
            <div class="transaction-details">
                <div class="detail-row">
                    <span class="detail-label">Reference</span>
                    <span class="detail-value">{{ reference }}</span>
                </div>
                
                <div class="detail-row">
                    <span class="detail-label">Date & Time</span>
                    <span class="detail-value">{{ transaction_time }}</span>
                </div>
                
                <div class="detail-row">
                    <span class="detail-label">Transaction ID</span>
                    <span class="detail-value">{{ transaction_id }}</span>
                </div>
                
                {% if narration %}
                <div class="detail-row">
                    <span class="detail-label">Description</span>
                    <span class="detail-value">{{ narration }}</span>
                </div>
                {% endif %}
            </div>
        </div>
        
        <div class="receipt-footer">
            <div class="qr-placeholder">QR Code</div>
            <div class="footer-text">
                Thank you for using Sofi AI!<br>
                For support, contact us via Telegram<br>
                Keep this receipt for your records
            </div>
        </div>
    </div>
</body>
</html>
        """
    
    def generate_html_receipt(self, transaction_data: Dict[str, Any]) -> str:
        """Generate HTML receipt from transaction data"""
        try:
            template = Template(self.template_html)
            
            # Format transaction time
            if isinstance(transaction_data.get('transaction_time'), str):
                transaction_time = transaction_data['transaction_time']
            else:
                transaction_time = datetime.now().strftime("%d/%m/%Y %I:%M %p")
            
            # Prepare template data
            template_data = {
                'amount': float(transaction_data.get('amount', 0)),
                'fee': float(transaction_data.get('fee', 0)),
                'total_charged': float(transaction_data.get('total_charged', 0)),
                'new_balance': float(transaction_data.get('new_balance', 0)),
                'recipient_name': transaction_data.get('recipient_name', 'Unknown Recipient'),
                'bank_name': transaction_data.get('bank_name', 'Unknown Bank'),
                'account_number': transaction_data.get('account_number', ''),
                'reference': transaction_data.get('reference', ''),
                'transaction_id': transaction_data.get('transaction_id', ''),
                'transaction_time': transaction_time,
                'narration': transaction_data.get('narration', '')
            }
            
            html_content = template.render(**template_data)
            return html_content
            
        except Exception as e:
            logger.error(f"Error generating HTML receipt: {e}")
            return None
    
    def generate_telegram_receipt(self, transaction_data: Dict[str, Any]) -> str:
        """Generate formatted receipt for Telegram"""
        try:
            amount = float(transaction_data.get('amount', 0))
            fee = float(transaction_data.get('fee', 0))
            total_charged = float(transaction_data.get('total_charged', 0))
            new_balance = float(transaction_data.get('new_balance', 0))
            
            receipt = f"""
🎉 *TRANSFER SUCCESSFUL!* 🎉

━━━━━━━━━━━━━━━━━━━━━━━━━
💳 *SOFI AI RECEIPT*
━━━━━━━━━━━━━━━━━━━━━━━━━

💰 *Amount Sent:* ₦{amount:,.2f}
💸 *Transfer Fee:* ₦{fee:,.2f}
💵 *Total Charged:* ₦{total_charged:,.2f}
💳 *New Balance:* ₦{new_balance:,.2f}

👤 *Recipient:* {transaction_data.get('recipient_name', 'Unknown')}
🏦 *Bank:* {transaction_data.get('bank_name', 'Unknown Bank')}
📱 *Account:* {transaction_data.get('account_number', '')}

🧾 *Reference:* `{transaction_data.get('reference', '')}`
🆔 *Transaction ID:* `{transaction_data.get('transaction_id', '')}`
🕐 *Time:* {transaction_data.get('transaction_time', datetime.now().strftime('%d/%m/%Y %I:%M %p'))}

━━━━━━━━━━━━━━━━━━━━━━━━━
Thank you for using Sofi AI! 💚
Keep this receipt for your records 📄
━━━━━━━━━━━━━━━━━━━━━━━━━
"""
            return receipt.strip()
            
        except Exception as e:
            logger.error(f"Error generating Telegram receipt: {e}")
            return "Receipt generation failed"
    
    def generate_pdf_receipt(self, transaction_data: Dict[str, Any], output_path: Optional[str] = None) -> Optional[str]:
        """Generate PDF receipt from HTML"""
        try:
            if not PDF_AVAILABLE:
                logger.warning("PDF generation not available - weasyprint not installed")
                return None
            
            # Generate HTML first
            html_content = self.generate_html_receipt(transaction_data)
            if not html_content:
                return None
            
            # Create output path if not provided
            if not output_path:
                timestamp = int(datetime.now().timestamp())
                output_path = f"/tmp/sofi_receipt_{timestamp}.pdf"
            
            # Generate PDF from HTML
            weasyprint.HTML(string=html_content).write_pdf(output_path)
            
            logger.info(f"PDF receipt generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating PDF receipt: {e}")
            return None
    
    def generate_word_doc_receipt(self, transaction_data: Dict[str, Any], output_path: Optional[str] = None) -> Optional[str]:
        """Generate Word document receipt"""
        try:
            # Try to import python-docx
            try:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                logger.warning("python-docx not available - Word document generation disabled")
                return None
            
            # Create document
            doc = Document()
            
            # Add header
            header = doc.add_heading('SOFI AI - TRANSACTION RECEIPT', 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add status
            status_para = doc.add_paragraph('✅ TRANSFER SUCCESSFUL')
            status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph('')  # Empty line
            
            # Add transaction details
            details_table = doc.add_table(rows=8, cols=2)
            details_table.style = 'Light Grid Accent 1'
            
            # Fill table with data
            details = [
                ('Amount Sent', f"₦{float(transaction_data.get('amount', 0)):,.2f}"),
                ('Transfer Fee', f"₦{float(transaction_data.get('fee', 0)):,.2f}"),
                ('Total Charged', f"₦{float(transaction_data.get('total_charged', 0)):,.2f}"),
                ('New Balance', f"₦{float(transaction_data.get('new_balance', 0)):,.2f}"),
                ('Recipient', transaction_data.get('recipient_name', 'Unknown')),
                ('Bank', transaction_data.get('bank_name', 'Unknown Bank')),
                ('Account Number', transaction_data.get('account_number', '')),
                ('Reference', transaction_data.get('reference', ''))
            ]
            
            for i, (label, value) in enumerate(details):
                details_table.cell(i, 0).text = label
                details_table.cell(i, 1).text = value
            
            # Add timestamp
            doc.add_paragraph('')
            time_para = doc.add_paragraph(f"Transaction Time: {transaction_data.get('transaction_time', datetime.now().strftime('%d/%m/%Y %I:%M %p'))}")
            time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add footer
            doc.add_paragraph('')
            footer = doc.add_paragraph('Thank you for using Sofi AI!')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            if not output_path:
                timestamp = int(datetime.now().timestamp())
                output_path = f"/tmp/sofi_receipt_{timestamp}.docx"
            
            doc.save(output_path)
            
            logger.info(f"Word document receipt generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating Word document receipt: {e}")
            return None

def create_transaction_receipt(transaction_data: Dict[str, Any], format_type: str = "telegram") -> str:
    """
    Create a transaction receipt
    
    Args:
        transaction_data: Transaction details
        format_type: "telegram", "html", or "pdf"
    
    Returns:
        Formatted receipt string
    """
    if format_type == "html":
        return receipt_generator.generate_html_receipt(transaction_data)
    elif format_type == "telegram":
        return receipt_generator.generate_telegram_receipt(transaction_data)
    elif format_type == "pdf":
        return receipt_generator.generate_pdf_receipt(transaction_data)
    elif format_type == "word":
        return receipt_generator.generate_word_doc_receipt(transaction_data)
    else:
        return receipt_generator.generate_telegram_receipt(transaction_data)
